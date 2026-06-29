import os
from docx import Document

class DocumentGenerator:
    def __init__(self):
        pass

    def save_narration_to_word(self, script_text, original_ppt_path):
        """
        Saves a structured backup copy of the AI lecture script outline into a Word file.
        """
        doc = Document()
        doc.add_heading('LectureForge AI — Generated Narrative Script', 0)
        
        # Populate content paragraphs
        doc.add_paragraph(script_text)
        
        # 🎯 BUILD SAFE CROSS-PLATFORM PATHS
        workspace_dir = os.path.dirname(original_ppt_path)
        base_name = os.path.basename(original_ppt_path).replace(".pptx", "")
        output_path = os.path.join(workspace_dir, f"Script_{base_name}.docx")
        
        doc.save(output_path)
        return output_path
